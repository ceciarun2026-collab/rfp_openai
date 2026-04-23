"""
report_generator.py
Genera el reporte Word de evaluación de propuestas.
Usa python-docx para construir el documento con formato profesional.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colores corporativos ──────────────────────────────────────────
COLOR_PRIMARIO = RGBColor(0x2E, 0x86, 0xC5)   # azul
COLOR_VERDE    = RGBColor(0x1E, 0x8B, 0x4C)   # habilitado
COLOR_ROJO     = RGBColor(0xC0, 0x39, 0x2B)   # no habilitado
COLOR_GRIS     = RGBColor(0x58, 0x58, 0x58)   # texto secundario


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_cell(cell, text: str, font_size: int = 11, color: RGBColor = None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = COLOR_PRIMARIO
    return p


def _add_info_row(table, label: str, value: str):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = COLOR_GRIS
    row.cells[1].text = str(value) if value else "N/D"


def generate_report(
    rfp_data: dict,
    scored_proposals: list,
    recommendation: dict,
    output_path: str = "output/reporte_evaluacion.docx"
) -> str:
    """
    Genera el reporte Word completo de evaluación.
    Retorna la ruta del archivo generado.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    proceso = rfp_data.get("proceso", {})
    fecha   = datetime.now().strftime("%d de %B de %Y")

    # ── 1. PORTADA ────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME DE EVALUACIÓN DE PROPUESTAS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_PRIMARIO

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(proceso.get("titulo", "Proceso de Contratación"))
    run2.font.size = Pt(14)
    run2.font.color.rgb = COLOR_GRIS

    doc.add_paragraph()
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_p.add_run(f"Generado el {fecha}").font.color.rgb = COLOR_GRIS

    doc.add_page_break()

    # ── 2. RESUMEN EJECUTIVO ───────────────────────────────────────
    _add_heading(doc, "1. Resumen Ejecutivo", level=1)

    habilitados    = [p for p in scored_proposals if p["habilitado"]]
    no_habilitados = [p for p in scored_proposals if not p["habilitado"]]

    exec_table = doc.add_table(rows=0, cols=2)
    exec_table.style = "Table Grid"
    exec_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    datos_resumen = [
        ("Entidad contratante",       proceso.get("entidad_contratante", "N/D")),
        ("Objeto del contrato",        proceso.get("objeto", "N/D")),
        ("Plazo de ejecución",         proceso.get("plazo_ejecucion", "N/D")),
        ("Total propuestas",           str(len(scored_proposals))),
        ("Propuestas habilitadas",     str(len(habilitados))),
        ("Propuestas no habilitadas",  str(len(no_habilitados))),
        ("Proveedor recomendado",      recommendation.get("ganador_sugerido", "N/D")),
        ("Fecha del informe",          fecha),
    ]
    for lbl, val in datos_resumen:
        _add_info_row(exec_table, lbl, val)

    doc.add_paragraph()
    if recommendation.get("recomendacion_final"):
        doc.add_paragraph(recommendation["recomendacion_final"])

    doc.add_page_break()

    # ── 3. CRITERIOS DEL RFP ──────────────────────────────────────
    _add_heading(doc, "2. Criterios del Proceso", level=1)

    _add_heading(doc, "2.1 Criterios Habilitantes", level=2)
    for c in rfp_data.get("criterios_habilitantes", []):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"[{c['id']}] {c['nombre']}: ")
        run.bold = True
        p.add_run(c.get("descripcion", ""))
        if c.get("valor_minimo"):
            doc.add_paragraph(
                f"      Valor mínimo requerido: {c['valor_minimo']}",
                style="List Bullet 2"
            )

    doc.add_paragraph()
    _add_heading(doc, "2.2 Criterios de Calificación", level=2)

    cal_table = doc.add_table(rows=1, cols=3)
    cal_table.style = "Table Grid"
    for i, h in enumerate(["Criterio", "Descripción", "Puntaje máximo"]):
        _bold_cell(cal_table.rows[0].cells[i], h, color=COLOR_PRIMARIO)
        _set_cell_bg(cal_table.rows[0].cells[i], "E8F0F8")

    for c in rfp_data.get("criterios_calificacion", []):
        row = cal_table.add_row()
        row.cells[0].text = f"[{c['id']}] {c['nombre']}"
        row.cells[1].text = c.get("descripcion", "")
        row.cells[2].text = str(c.get("puntaje_maximo", 0))
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Puntaje total: {rfp_data.get('puntaje_total', 0)} puntos")
    run.bold = True

    doc.add_page_break()

    # ── 4. EVALUACIÓN HABILITANTES ────────────────────────────────
    _add_heading(doc, "3. Evaluación de Criterios Habilitantes", level=1)

    criterios_hab = rfp_data.get("criterios_habilitantes", [])
    cols = ["Proveedor"] + [c["nombre"] for c in criterios_hab] + ["¿Habilitado?"]
    hab_table = doc.add_table(rows=1, cols=len(cols))
    hab_table.style = "Table Grid"

    for i, h in enumerate(cols):
        _bold_cell(hab_table.rows[0].cells[i], h, font_size=10, color=COLOR_PRIMARIO)
        _set_cell_bg(hab_table.rows[0].cells[i], "E8F0F8")

    for proposal in scored_proposals:
        row = hab_table.add_row()
        row.cells[0].text = proposal["provider"]

        hab_detail  = proposal.get("habilitacion_detalle", {})
        criterios_ev = hab_detail.get("criterios_habilitantes", [])

        for j, criterio_rfp in enumerate(criterios_hab, 1):
            cid   = criterio_rfp["id"]
            match = next((c for c in criterios_ev if c.get("id") == cid), None)
            if match:
                cumple = match.get("cumple", False)
                cell   = row.cells[j]
                cell.text = "✓" if cumple else "✗"
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cell.paragraphs[0].runs[0]
                r.font.color.rgb = COLOR_VERDE if cumple else COLOR_ROJO
                r.bold = True
            else:
                row.cells[j].text = "–"

        es_hab = proposal["habilitado"]
        cell   = row.cells[-1]
        cell.text = "SÍ" if es_hab else "NO"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].runs[0]
        r.bold = True
        r.font.color.rgb = COLOR_VERDE if es_hab else COLOR_ROJO

    doc.add_page_break()

    # ── 5. TABLA DE PUNTAJES ──────────────────────────────────────
    _add_heading(doc, "4. Calificación y Ranking", level=1)

    criterios_cal = rfp_data.get("criterios_calificacion", [])
    col_names = (["#", "Proveedor", "Precio ofertado"]
                 + [c["nombre"] for c in criterios_cal]
                 + ["Puntaje total", "%"])

    score_table = doc.add_table(rows=1, cols=len(col_names))
    score_table.style = "Table Grid"

    for i, h in enumerate(col_names):
        _bold_cell(score_table.rows[0].cells[i], h, font_size=10, color=COLOR_PRIMARIO)
        _set_cell_bg(score_table.rows[0].cells[i], "E8F0F8")

    for proposal in habilitados:
        row = score_table.add_row()
        row.cells[0].text = str(proposal["ranking"])
        row.cells[1].text = proposal["provider"]
        precio = proposal.get("precio_ofertado")
        row.cells[2].text = f"${precio:,.0f}" if precio else "N/D"

        desglose = {d["id"]: d for d in proposal.get("desglose_puntaje", [])}
        for j, criterio in enumerate(criterios_cal, 3):
            d   = desglose.get(criterio["id"], {})
            pts = d.get("puntaje_obtenido", 0)
            row.cells[j].text = f"{pts:.2f}"
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        row.cells[-2].text = f"{proposal['puntaje_total']:.2f}"
        row.cells[-2].paragraphs[0].runs[0].bold = True
        row.cells[-1].text = f"{proposal['porcentaje']:.1f}%"

    for proposal in no_habilitados:
     row = score_table.add_row()

    # Columna # y Proveedor
    row.cells[0].text = "N/H"
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = COLOR_ROJO

    row.cells[1].text = proposal["provider"]

    # Precio
    precio = proposal.get("precio_ofertado")
    row.cells[2].text = f"${precio:,.0f}" if precio else "N/D"

    # ✅ FIX: escribir los puntajes por criterio (con fondo rojo claro)
    desglose = {d["id"]: d for d in proposal.get("desglose_puntaje", [])}
    for j, criterio in enumerate(criterios_cal, 3):
        d = desglose.get(criterio["id"], {})
        pts = d.get("puntaje_obtenido", 0)
        cell = row.cells[j]
        _set_cell_bg(cell, "FDECEA")  # fondo rojo claro
        cell.text = f"{pts:.2f}" if pts is not None else "—"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ✅ FIX: mostrar puntaje total y % reales en lugar de "N/H"
    cell_total = row.cells[-2]
    _set_cell_bg(cell_total, "FDECEA")
    cell_total.text = f"{proposal['puntaje_total']:.2f}"
    cell_total.paragraphs[0].runs[0].bold = True
    cell_total.paragraphs[0].runs[0].font.color.rgb = COLOR_ROJO

    cell_pct = row.cells[-1]
    _set_cell_bg(cell_pct, "FDECEA")
    cell_pct.text = f"{proposal['porcentaje']:.1f}%"
    cell_pct.paragraphs[0].runs[0].font.color.rgb = COLOR_ROJO

    doc.add_page_break()

    # ── 6. DETALLE POR PROVEEDOR ──────────────────────────────────
    _add_heading(doc, "5. Detalle por Proveedor", level=1)

    for proposal in scored_proposals:
        _add_heading(doc, f"5. {proposal['provider']}", level=2)

        estado       = "HABILITADO" if proposal["habilitado"] else "NO HABILITADO"
        color_estado = COLOR_VERDE if proposal["habilitado"] else COLOR_ROJO

        p   = doc.add_paragraph()
        run = p.add_run(f"Estado: {estado}")
        run.bold = True
        run.font.color.rgb = color_estado

        hab = proposal.get("habilitacion_detalle", {})
        if hab.get("criterios_habilitantes"):
            doc.add_paragraph("Criterios habilitantes:", style="Intense Quote")
            for c in hab["criterios_habilitantes"]:
                bullet = doc.add_paragraph(style="List Bullet")
                mark   = "✓" if c.get("cumple") else "✗"
                run    = bullet.add_run(f"{mark} {c['nombre']}: ")
                run.bold = True
                run.font.color.rgb = COLOR_VERDE if c.get("cumple") else COLOR_ROJO
                bullet.add_run(c.get("evidencia", "Sin evidencia encontrada"))

        if proposal.get("desglose_puntaje"):
            doc.add_paragraph()
            doc.add_paragraph("Puntaje por criterio:", style="Intense Quote")
            for d in proposal["desglose_puntaje"]:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(f"{d['nombre']}: ").bold = True
                bullet.add_run(f"{d['puntaje_obtenido']:.2f} / {d['puntaje_maximo']} pts")
                if d.get("evidencia"):
                    doc.add_paragraph(
                        f"   Evidencia: {d['evidencia']}",
                        style="List Bullet 2"
                    )

        doc.add_paragraph()

    doc.add_page_break()

    # ── 7. RECOMENDACIÓN ─────────────────────────────────────────
    _add_heading(doc, "6. Recomendación de Adjudicación", level=1)

    if recommendation.get("ganador_sugerido"):
        p   = doc.add_paragraph()
        run = p.add_run(f"Proveedor recomendado: {recommendation['ganador_sugerido']}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_PRIMARIO

    if recommendation.get("justificacion"):
        doc.add_paragraph()
        _add_heading(doc, "Justificación", level=2)
        doc.add_paragraph(recommendation["justificacion"])

    if recommendation.get("analisis_comparativo"):
        doc.add_paragraph()
        _add_heading(doc, "Análisis comparativo", level=2)
        doc.add_paragraph(recommendation["analisis_comparativo"])

    if recommendation.get("alertas"):
        doc.add_paragraph()
        _add_heading(doc, "Alertas y consideraciones", level=2)
        for alerta in recommendation["alertas"]:
            doc.add_paragraph(alerta, style="List Bullet")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"[OK] Reporte guardado: {output_path}")
    return output_path
